#!/usr/bin/env python3
#
# Copyright (c) 2019, Arista Networks EOS+
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# * Redistributions of source code must retain the above copyright notice, this
#   list of conditions and the following disclaimer.
#
# * Redistributions in binary form must reproduce the above copyright notice,
#   this list of conditions and the following disclaimer in the documentation
#   and/or other materials provided with the distribution.
#
# * Neither the name of the Arista nor the names of its
#   contributors may be used to endorse or promote products derived from
#   this software without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE AR
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE
# FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL
# DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
# SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER
# CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
# OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
# OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#

""" Utility script for taking a manual testing spreadsheet and converting it
    it into a formattted word doc.

    Input Variables:
        CSVFILES (list): List of path + file names for CSV input data
        DOCFILE (str): Path + file name for output Word doc
        EOS (str): EOS version number
        TITLE_ROW (int): Title row in CSV
        HEADER_ROW (int): Header row in CSV
        TC_ID (str): CSV column name with test case ID
        CASE_NAME (str): CSV column name with test case name
        PROCEDURE (str): CSV column name with test steps / procedures
        EXPECTED_RESULT (str): CSV column name with test criteria
        PASS_FAIL (str): CSV column name with test status: pass or fail
        OBSERVATION (str): CSV column name with test completion note
        TC_TYPE (str): CSV column name with test case type
"""

import csv
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_COLOR_INDEX

# Global Variables
# test suite header
TEST_SUITE_HEADER = r"\d{1,2}\.\d{1,2}"
# Input Variables
# Name of CSV to use for test output
CSVFILES = [
    "../Arista Agora Test Plan V9.3 - EOS Test Items.csv",
    "../Arista Agora Test Plan V9.3 - CVA_CVP Test Items.csv",
]
DOCFILE = "../test_summary_results.docx"
# EOS version
EOS = "4.29.1F-DPE"
# Column name of interesting values
TITLE_ROW = 0
HEADER_ROW = 1
TC_ID = r".*Test Case ID.*"
CASE_NAME = r".*Test Case Description.*"
PROCEDURE = r".*Test Steps.*"
EXPECTED_RESULT = r".*Criteria.*"
PASS_FAIL = r"Pass/Fail"
OBSERVATION = r".*Observations.*"
TC_TYPE = r".*Test Case Type.*"
# Column Width
WIDTHS = (
    Inches(0.5),
    Inches(1.17),
    Inches(2.50),
    Inches(1.25),
    Inches(0.75),
    Inches(1.50),
)


def parse_csv():
    """Iterate through a list of CSV files.  Open and parse CSV data
       and find test data

    Returns:
        list: List of interesting test data
    """
    test_data = []

    # iterate through CSV files
    for csvfile in CSVFILES:
        print(f"Parsing CSV file: {csvfile} for test data...")

        # open csv
        with open(csvfile, "r", encoding="utf-8") as file:
            csv_reader = csv.reader(file)
            # regex for matching a test suite header
            inspect_csv_data(test_data, csv_reader)

    return test_data


def set_header_columns(row):
    """Use regex to find a column index for interesting data

    Args:
        row (list): CSV row

    Returns:
        dict: key is column name and value is column index
    """

    header_columns = {}

    # iterate through cells to find column headers
    for count, _ in enumerate(row):
        if re.match(TC_ID, row[count]):
            header_columns["tc_id"] = count
        elif re.match(CASE_NAME, row[count]):
            header_columns["case_name"] = count
        elif re.match(PROCEDURE, row[count]):
            header_columns["procedure"] = count
        elif re.match(EXPECTED_RESULT, row[count]):
            header_columns["expected_result"] = count
        elif re.match(PASS_FAIL, row[count]):
            header_columns["pass_fail"] = count
        elif re.match(OBSERVATION, row[count]):
            header_columns["observation"] = count
        elif re.match(TC_TYPE, row[count]):
            header_columns["tc_type"] = count

    return header_columns


def inspect_csv_data(test_data, csv_reader):
    """Inspect each row of CSV data and determine if its interesting

    Args:
        test_data (list): Use CSV file to update list with interesting test data
        csv_reader (obj): CSV object to iterate through
    """

    for count, row in enumerate(csv_reader):
        test_data_entry = []

        # skip title row
        if count == TITLE_ROW:
            pass
        # inspect header and find correct column numbers
        elif count == HEADER_ROW:
            header_columns = set_header_columns(row)
        # parse csv row for a test case ids
        elif "TN" in row[header_columns["tc_id"]]:
            test_data_entry.append(row[header_columns["tc_id"]])
            test_data_entry.append(row[header_columns["case_name"]])
            test_data_entry.append(row[header_columns["procedure"]])
            test_data_entry.append(row[header_columns["expected_result"]])
            test_data_entry.append(row[header_columns["pass_fail"]])
            test_data_entry.append(row[header_columns["observation"]])
            test_data.append(test_data_entry)
        # parse csv row for header
        elif re.match(TEST_SUITE_HEADER, row[header_columns["case_name"]]):
            test_data_entry.append(row[header_columns["case_name"]])
            test_data.append(test_data_entry)
        # parse csv row for sub-section
        elif (not row[header_columns["tc_id"]] and not row[header_columns["tc_type"]]) and row[
            header_columns["case_name"]
        ] != "":
            test_data_entry.append("")
            test_data_entry.append(row[header_columns["case_name"]])
            test_data.append(test_data_entry)


def write_report(test_data):
    """Write word doc with test data collected from CSV

    Args:
        test_data (list): List of interesting test data
    """
    print(f"Writing report doc: {DOCFILE}")

    # Open doc
    document = docx.Document()
    # Change doc margins
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    write_header(document)
    write_results_table(document, test_data)

    doc_text = (
        "\n\n\nPlease note – for more details on test outputs please "
        "reference the test case output appendix docs.\n"
    )
    write_text(document, doc_text)

    # Close doc
    document.save(DOCFILE)


def write_header(document):
    """Write test section header to word doc

    Args:
        document (obj): Test report Word doc
        test_data (list): List of interesting test data
    """

    # Write summary header
    heading = document.add_heading(level=2)
    run = heading.add_run("4	Test Summary Result")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(204, 0, 0)

    # Header text
    doc_text = (
        f"\nTest Cases for Arista BAU Hardware & EOS {EOS} Certification. "
        "Please note – for more details on test outputs please reference "
        "the test case output appendix docs.\n\n\n"
    )
    write_text(document, doc_text)


def write_text(document, doc_text, font_pt=10, font="Arial"):
    """Write a paragraph to Word Doc

    Args:
        document (obj): Word doc object
        doc_text (str): Text to write to doc
        pt (int): Word font size
        font (string): Word font style
    """
    para = document.add_paragraph()
    run = para.add_run(doc_text)
    run.font.size = Pt(font_pt)
    run.font.name = font


def write_header_row(table):
    """Write the header row for the result table

    Args:
        table (obj): Word doc table object
    """

    headers = [
        "Test Case",
        "Case Name",
        "Procedure",
        "Expected result",
        "Pass/Fail",
        "Observation",
    ]

    # Write header row in table
    for count, header in enumerate(headers):
        set_cell_widths(table, 0)
        para = table.rows[0].cells[count].paragraphs[0]

        run = para.add_run(header)
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.bold = True

        shade_cell(count, table, 0, "00FFFF")


def write_results_table(document, test_data):
    """Write results table to Word doc

    Args:
        document (obj): Word doc object
        test_data (list): List of interesting test data
    """

    # create table
    table = document.add_table(rows=1, cols=6, style="Table Grid")
    table.autofit = False

    # create table rows
    write_header_row(table)
    write_data_row(table, test_data)


def write_data_row(table, test_data):
    """Write a test case result row in table

    Args:
        table (obj): Word doc table object
        test_data (list): List of interesting test data
    """

    # Iterate through test data and write table rows
    for count, test_case_entry in enumerate(test_data):
        count += 1
        row_cells = table.add_row().cells
        set_cell_widths(table, count)

        # Identify test suite sub-header row
        if re.match(TEST_SUITE_HEADER, test_case_entry[0]):
            row_cells[0].merge(row_cells[5])
            row_cells[0].text = ""
            run = row_cells[0].paragraphs[0].add_run(test_case_entry)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.bold = True
            shade_cell(0, table, count, "CCFFFF")

        # Identify test case result data row
        elif "TN" in test_case_entry[0]:
            for counter, test_data_entry in enumerate(test_case_entry):
                row_cells[counter].text = ""
                run.font.name = "Arial"
                run.font.size = Pt(9)

                # Test Case Identifier cell
                if counter == 0:
                    test_data_entry = format_test_id(test_data_entry)
                    run = row_cells[counter].paragraphs[0].add_run(test_data_entry)
                # Test Case Pass/Fail cell
                elif counter == 4:
                    format_pass_fail(row_cells, counter, test_data_entry)
                # Cells not needing special formatting
                else:
                    run = row_cells[counter].paragraphs[0].add_run(test_data_entry)

                run.font.name = "Arial"
                run.font.size = Pt(9)

        # identify test suite sub section
        elif test_case_entry[0] == "":
            row_cells[2].merge(row_cells[5])

            run = row_cells[1].paragraphs[0].add_run(test_case_entry[1])
            run.font.name = "Arial"
            run.font.size = Pt(9)

            for cell_idx in [0, 1, 2]:
                shade_cell(cell_idx, table, count, "CCFFFF")


def shade_cell(cell_idx, table, count, shade):
    """Shade a cell in word doc table

    Args:
        cell_idx (int): Column index for cell to shade
        table (obj): Word doc table object
        count (int): Row index for cell to shade
        shade (str): hexadecimal color representation
    """
    cell = table.cell(count, cell_idx)
    color = parse_xml(
        # pylint: disable-next=consider-using-f-string
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), shade)
    )
    # pylint: disable-next=protected-access
    cell._tc.get_or_add_tcPr().append(color)


def format_test_id(test_data_entry):
    """Format the test case id

    Args:
        test_data_entry (str): Text for table cell

    Returns:
        str: Formatted text for table cell
    """
    if "TN" in test_data_entry:
        test_data_entry = test_data_entry.split("TN")[1]

    return test_data_entry


def format_pass_fail(row_cells, counter, test_data_entry):
    """Format pass fail cell

    Args:
        row_cells (obj): Table cell
        counter (int): Table cell number
        test_data_entry (str): Text for table cell
    """
    run = row_cells[counter].paragraphs[0].add_run(test_data_entry.upper())

    if test_data_entry.upper() == "PASS":
        # pylint: disable-next=no-member
        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        run.font.bold = True
    elif test_data_entry.upper() == "FAIL":
        # pylint: disable-next=no-member
        run.font.highlight_color = WD_COLOR_INDEX.RED
        run.font.bold = True
    else:
        # pylint: disable-next=no-member
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.bold = True


def set_cell_widths(table, row):
    """Set width of Word table cell

    Args:
        table (obj): Word doc table object
        row (int): Table row number for operations
    """

    for column, width in enumerate(WIDTHS):
        table.rows[row].cells[column].width = width


def main():
    """Main Python function"""

    print("Starting...")
    test_data = parse_csv()
    write_report(test_data)
    print("Complete")


if __name__ == "__main__":
    main()
