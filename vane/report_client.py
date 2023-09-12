#!/usr/bin/env python3
#
# Copyright (c) 2023, Arista Networks EOS+
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# * Redistributions of source code must retain the above copyright notice, this
#   list of conditions and the following disclaimer.
#
# * Redistributions in binary form must reproduce the above copyright notice,
#   this list of conditions and the following disclaimer in the documentation
#   and/or other materials provided with the distribution.
#
# * Neither the name of the Arista nor the names of its
#   contributors may be used to endorse or promote products derived from
#   this software without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE
# FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL
# DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
# SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER
# CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
# OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
# OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#
# pylint: disable=too-many-lines

"""Utilities for using PyTest in network testing"""

import json
import os
import re
import yaml
import docx
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Inches, Pt, RGBColor
from vane.report_templates import REPORT_TEMPLATES
from vane.tests_tools import yaml_read
from vane.vane_logging import logging
from vane.utils import return_date

TABLE_GRID = "Table Grid"
TOTAL_TESTS = "Total Tests"
TOTAL_PASSED = "Total Passed"
TOTAL_FAILED = "Total Failed"
TOTAL_SKIPPED = "Total Skipped"


# pylint: disable=too-few-public-methods
class ReportClient:
    """Creates an instance of the Report Client."""

    def __init__(self, test_definition):
        """Initializes the Report Client

        Args:
            test_definition (str): YAML representation of NRFU tests
        """

        logging.info("Reading YAML data-model and converting into a Python data structure")
        self.data_model = yaml_read(test_definition)
        logging.debug(f"Internal test data-model initialized with value: {self.data_model}")
        self._summary_results = self._compile_test_results()
        logging.debug(f"Test Results: {self._summary_results}")

        self._reports_dir = self.data_model["parameters"]["report_dir"]
        _results_dir = self.data_model["parameters"]["results_dir"]
        self._results_datamodel = None
        self._compile_yaml_data(_results_dir)
        logging.debug(f"Results file data is {self._results_datamodel}")

        self._document = docx.Document()
        section = self._document.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        self._major_section = 1
        self._test_no = 1

    def _compile_yaml_data(self, yaml_dir):
        """Import result data into report client to results data model

        Args:
            yaml_dir (str): Results directory path
        """

        logging.info("Compiling test case results into data model")
        logging.debug(f"yaml directory is {yaml_dir}\n")
        yaml_files = os.listdir(yaml_dir)
        logging.debug(f"yaml input files are {yaml_files}")

        for name in yaml_files:
            if "result-" in name:
                yaml_file = f"{yaml_dir}/{name}"
                yaml_data = yaml_read(yaml_file)

                self._reconcile_results(yaml_data)
            else:
                logging.error(f"Incorrect filename: {name}")

        logging.debug(f"Updated results_data to {self._results_datamodel}")

    def _reconcile_results(self, test_parameters):
        """Validate test case results data and reconciles any missing data

        Args:
            test_parameters (dict): data struct representing a test case
        """

        test_suite = test_parameters["test_suite"]
        test_suite = test_suite.split("/")[-1]
        dut_name = test_parameters["dut"]
        test_case = test_parameters["name"]
        logging.info(
            "Validating test case results data and reconciling missing data for test case: "
            f"{test_case} on DUT: {dut_name}"
        )

        if not self._results_datamodel:
            self._results_datamodel = {
                "test_suites": [
                    {
                        "name": test_suite,
                        "test_cases": [{"name": test_case, "duts": []}],
                    }
                ]
            }

        logging.debug(f"\n\n\rFound the Index for test suite: {test_suite} on dut {dut_name}\n\n\r")
        logging.debug(self._results_datamodel["test_suites"])
        test_suites = [param["name"] for param in self._results_datamodel["test_suites"]]

        if test_suite in test_suites:
            suite_index = test_suites.index(test_suite)
            logging.debug(f"Test suite {test_suite} exists in results file at index {suite_index}")
        else:
            logging.info(f"Create test suite {test_suite} in results file")
            suite_stub = {"name": test_suite, "test_cases": []}
            self._results_datamodel["test_suites"].append(suite_stub)
            suite_index = len(self._results_datamodel["test_suites"]) - 1

        test_cases = [
            param["name"]
            for param in self._results_datamodel["test_suites"][suite_index]["test_cases"]
        ]

        if test_case in test_cases:
            test_index = test_cases.index(test_case)
            logging.debug(f"Test case {test_case} exists in results file at index {test_index}")
        else:
            logging.info(f"Create test case {test_case} in results file")
            test_stub = {"name": test_case, "duts": []}
            self._results_datamodel["test_suites"][suite_index]["test_cases"].append(test_stub)
            test_index = len(self._results_datamodel["test_suites"][suite_index]["test_cases"]) - 1

        logging.debug(f"Find Index for dut {dut_name}")
        duts = [
            param["dut"]
            for param in self._results_datamodel["test_suites"][suite_index]["test_cases"][
                test_index
            ]["duts"]
        ]

        if dut_name not in duts:
            logging.debug(
                f"Add DUT {dut_name} to test case {test_case} with parameters {test_parameters}"
            )
            yaml_ptr = self._results_datamodel["test_suites"][suite_index]
            yaml_ptr["test_cases"][test_index]["duts"].append(test_parameters)

    def write_result_doc(self):
        """Create MSFT docx with results"""

        logging.info("Create MSFT docx with results")
        self._write_title_page()
        self._write_toc_page()
        self._write_summary_report()
        self._write_tests_case_report()
        self._write_detail_report()

        _, file_date = return_date()
        reports_dir = self._reports_dir
        file_name = f"{reports_dir}/report_{file_date}.docx"
        logging.info(f"Writing docx report to file: {file_name}")
        self._document.save(file_name)

    def _write_title_page(self):
        """Write report title page"""

        logging.info("Creating report title page")
        format_date, _ = return_date()
        self._document.add_heading("Test Report", 0)
        para = self._document.add_paragraph()
        self._write_text(para, format_date, align="right")
        self._document.add_page_break()

    def _write_toc_page(self):
        """Write table of contents page"""

        logging.info("Creating report table of contents page")
        w_fldchar = "w:fldChar"
        w_fldchar_type = "w:fldCharType"

        paragraph = self._document.add_paragraph()
        run = paragraph.add_run()
        fld_char = OxmlElement(w_fldchar)
        fld_char.set(qn(w_fldchar_type), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

        fld_char2 = OxmlElement(w_fldchar)
        fld_char2.set(qn(w_fldchar_type), "separate")
        fld_char3 = OxmlElement("w:t")
        fld_char3.text = "Right-click to update field."
        fld_char2.append(fld_char3)

        fld_char4 = OxmlElement(w_fldchar)
        fld_char4.set(qn(w_fldchar_type), "end")

        # pylint: disable-next=protected-access
        r_element = run._r
        r_element.append(fld_char)
        r_element.append(instr_text)
        r_element.append(fld_char2)
        r_element.append(fld_char4)

        self._document.add_page_break()

    def _write_summary_report(self):
        """Write summary reports"""

        logging.info("Creating summary report section")
        self._document.add_heading(f"{self._major_section}. Test Results Summary", 1)
        self._write_summary_results()
        self._write_dut_summary_results()
        self._write_suite_summary_results()

        self._major_section += 1
        self._document.add_page_break()

    def _write_summary_results(self):
        """Write summary results section"""

        logging.info("Creating total test case summary results table")
        self._document.add_heading(f"{self._major_section}.1 Summary Results", 2)
        table = self._document.add_table(rows=1, cols=6, style="Table Grid")
        headers = [
            TOTAL_TESTS,
            TOTAL_PASSED,
            TOTAL_FAILED,
            TOTAL_SKIPPED,
            "Total Errored",
            "Total Duration",
        ]

        for column, header in enumerate(headers):
            self._write_cell(table, header.upper(), column, 0, "Arial", 9, True, "00FFFF")

        _ = table.add_row().cells
        data_row = []
        ptr = self._summary_results["summaryResults"]
        data_row.append(self._totals(ptr, "num_tests"))
        data_row.append(self._totals(ptr, "passed"))
        data_row.append(self._totals(ptr, "failed"))
        data_row.append(self._totals(ptr, "skipped"))
        data_row.append(self._totals(ptr, "error"))
        data_row.append(self._totals(ptr, "duration"))

        for column, cell_data in enumerate(data_row):
            self._write_cell(table, cell_data, column, 1, "Arial", 9)

    def _write_dut_summary_results(self):
        """Write summary DUT result section"""

        logging.info("Creating DUT summary results table")
        self._document.add_heading(
            f"{self._major_section }.2 Summary Totals for Devices Under Tests",
            2,
        )

        table = self._document.add_table(rows=1, cols=6, style="Table Grid")
        headers = ["DUT", TOTAL_TESTS, TOTAL_PASSED, TOTAL_FAILED, TOTAL_SKIPPED, "Total Errored"]

        for column, header in enumerate(headers):
            self._write_cell(table, header.upper(), column, 0, "Arial", 9, True, "00FFFF")

        duts = self._summary_results["duts"]

        for row, dut in enumerate(duts):
            logging.debug(f"Creating dut summary row: {row+1}")
            _ = table.add_row().cells
            data_row = []

            data_row.append(self._totals(dut, "name"))
            data_row.append(self._totals(dut, "TOTAL"))
            data_row.append(self._totals(dut, "PASS"))
            data_row.append(self._totals(dut, "FAIL"))
            data_row.append(self._totals(dut, "SKIP"))
            data_row.append(self._totals(dut, "ERROR"))

            for column, cell_data in enumerate(data_row):
                self._write_cell(table, cell_data, column, (row + 1), "Arial", 9)

    def _write_suite_summary_results(self):
        """Write summary test suite result section"""

        logging.info("Create Suite summary results table")
        self._document.add_heading(f"{self._major_section }.3 Summary Totals for Test Suites", 2)
        suite_results = self._compile_suite_results()
        if not suite_results:
            logging.warning("Skipping the test suite results")
            return

        table = self._document.add_table(rows=1, cols=5, style="Table Grid")
        headers = ["Test Suite", TOTAL_TESTS, TOTAL_PASSED, TOTAL_FAILED, TOTAL_SKIPPED]

        for column, header in enumerate(headers):
            self._write_cell(table, header.upper(), column, 0, "Arial", 9, True, "00FFFF")

        if not suite_results:
            logging.warning("Skipping the test suite results")
            return

        for row, suite_result in enumerate(suite_results):
            _ = table.add_row().cells
            data_row = []
            ts_name = self._format_ts_name(suite_result["name"])
            logging.debug(f"Writing row {row+1}")

            data_row.append(ts_name)
            data_row.append(str(suite_result["total_tests"]))
            data_row.append(str(suite_result["total_pass"]))
            data_row.append(str(suite_result["total_fail"]))
            data_row.append(str(suite_result["total_skip"]))

            for column, cell_data in enumerate(data_row):
                self._write_cell(table, cell_data, column, (row + 1), "Arial", 9)

    def _compile_test_results(self):
        """Parse PyTest JSON results and compile:"""

        json_report = self.data_model["parameters"]["json_report"]
        json_report = f"{json_report}.json"
        test_results = {}
        logging.info(f"Opening JSON file {json_report} to parse for summary results")

        with open(json_report, "r", encoding="utf-8") as json_file:
            logging.debug(f"Raw json report is {json_file}")
            test_data = json.load(json_file)
            tests = test_data["report"]["tests"]
            logging.debug(f"Structured json report is {test_data}")

            summary = test_data["report"]["summary"]
            test_results["summaryResults"] = summary
            logging.debug(f"Summary for test cases are {summary}")
            test_results["duts"] = self._parse_testcases(tests)

        return test_results

    def _parse_testcases(self, testcases):
        """Parse Test cases and return compilation per DUT

        Args:
            testcases (dict): data structure of test case results

        Returns:
            dict: data structure with compiled results
        """
        testcases_results = []
        dut_list = []

        for testcase in testcases:
            if re.search(r"\[.*\]", testcase["name"]):
                dut_name = re.findall(r"\[.*\]", testcase["name"])[0][1:-1]
                test_result = testcase["outcome"]

                if dut_name not in dut_list:
                    dut_list.append(dut_name)
                    testcases_results.append({})
                    testcases_results[-1]["PASS"] = 0
                    testcases_results[-1]["FAIL"] = 0
                    testcases_results[-1]["SKIP"] = 0
                    testcases_results[-1]["ERROR"] = 0
                    testcases_results[-1]["TOTAL"] = 0

                dut_index = dut_list.index(dut_name)
                testcases_results[dut_index]["name"] = dut_name

                if test_result == "passed":
                    testcases_results[dut_index]["PASS"] += 1
                elif test_result == "failed":
                    testcases_results[dut_index]["FAIL"] += 1
                elif test_result == "skipped":
                    testcases_results[dut_index]["SKIP"] += 1
                elif test_result == "error":
                    testcases_results[dut_index]["ERROR"] += 1

                testcases_results[dut_index]["TOTAL"] += 1

        logging.debug(f"DUT compiled results: {testcases_results}")

        return testcases_results

    def _totals(self, ptr, ptr_key):
        """Test for a key in dictionary.  If key exists return key and if key is
            missing return 0

        Args:
            ptr (dict): dictionary to check
            ptr_key (str): key to test if in dict

        Retrun:
            total (str): Value in dictionary
        """

        total = "0"

        if ptr_key in ptr:
            total = str(ptr[ptr_key])

        return total

    def _write_tests_case_report(self):
        """Write summary test case report"""

        logging.info("Creating report summary test case report section")
        self._document.add_heading(f"{self._major_section}. Test Case Results Summary", 1)
        self._major_section += 1

        report_styles = REPORT_TEMPLATES.keys()
        logging.debug(f"Inputted the following report summary styles {report_styles}")

        if "report_summary_style" in self.data_model["parameters"]:
            report_style = self.data_model["parameters"]["report_summary_style"]
            logging.debug(f"Summary style in parameters set to {report_style}")
            if report_style in report_styles:
                logging.debug(f"report_summary_style is set to {report_style}, custom")
                report_template = REPORT_TEMPLATES[report_style]
                self._custom_tc_report(report_template)
            else:
                logging.debug(f"report_summary_style is correctly set to {report_style}, default")
                self._default_tc_report()
        else:
            logging.warning(f"No summary style set in parameters: {self.data_model['parameters']}")
            self._default_tc_report()

    def _custom_tc_report(self, report_template):
        """Writes a custom test case summary report based on report template.
           Report template field must have summary: True in order for field to be displayed
           in report.

        Args:
            report_template (dict): Data structure describing reports fields
        """
        summary_headers = self._return_summary_headers(report_template)
        testcase_results = self._compile_custom_tc_results(summary_headers)
        self._write_custom_tc_report(summary_headers, testcase_results, report_template)

    def _write_custom_tc_report(self, summary_headers, testcase_results, report_template):
        """Writes custom test case report to Word doc

        Args:
            summary_headers (list): List of fields to include in test case report header row
            testcase_results (dict): Data structure with test case results
            report_template (dict): Data structure describing reports fields
        """
        columns = len(summary_headers)
        table = self._document.add_table(rows=1, cols=columns, style="Table Grid")

        if testcase_results:
            self._create_header_row(table, summary_headers, report_template)
            self._create_data_row(table, testcase_results, report_template)
        else:
            logging.error("Vane ran no test cases against DUTS")
            print(
                "Vane ran no test cases against DUTs.  Check your duts.yml file and test case "
                "filters"
            )

    def _create_header_row(self, table, summary_headers, report_template):
        """Writes header row within Word doc table

        Args:
            table (obj): Word doc obj representing a tables
            testcase_results (dict): Data structure with test case results
            report_template (dict): Data structure describing reports fields
        """
        headers = []
        row = 0

        for summary_header in summary_headers:
            logging.debug(f"summary header info: {summary_header}")
            if "output_name" in report_template[summary_header]:
                logging.debug(
                    f"Output name set to: {report_template[summary_header]['output_name']}"
                )
                headers.append(report_template[summary_header]["output_name"])
            else:
                logging.warning(f"No output name setting header to: {summary_header}")
                headers.append(summary_header)

        for column, header in enumerate(headers):
            self._write_cell(table, header.upper(), column, row, "Arial", 9, True, "00FFFF")

    def _create_data_row(self, table, testcase_results, report_template):
        """Writes a data row within Word doc table

        Args:
            table (obj): Word doc obj representing a tables
            testcase_results (dict): Data structure with test case results
            report_template (dict): Data structure describing reports fields
        """

        for row, testcase_result in enumerate(testcase_results):
            _ = table.add_row().cells
            for column, testcase_data in enumerate(testcase_result):
                logging.debug(
                    f"Writing test field: {testcase_data}"
                    f"with value: {report_template[testcase_data]}"
                )
                if "format" in report_template[testcase_data]:
                    data_format = report_template[testcase_data]["format"]
                    logging.debug(f"Format has been set to {format}")

                self._write_cell(
                    table,
                    testcase_result[testcase_data],
                    column,
                    (row + 1),
                    "Arial",
                    9,
                    data_format=data_format,
                )

    # pylint: disable-next=too-many-arguments
    def _write_cell(
        self,
        table,
        text,
        column,
        row,
        font=None,
        font_size=None,
        bold=False,
        color=None,
        data_format="string",
        text_color=None,
    ):
        """Writes a cell within Word doc table

        Args:
            table (obj): Word doc obj representing a table
            text (str): Text to output in table cell
            column (int): Column number in table cell
            row (int): Row number in table cell
            font (str, optional): Font to use in table cell. Defaults to None.
            font_size (int, optional): Font size to use in table cell. Defaults to None.
            bold (bool, optional): Bold text in table cell. Defaults to False.
            color (str, optional): Hex-decimanal color to fill table cell. Defaults to None.
            format (str, optional): Style of outputting text in table cell. Defaults to "string".
            text_color (obj, optional): Text output color. Defaults to None
        """
        para = table.rows[row].cells[column].paragraphs[0]
        logging.debug(f"Added cell ({row}, {column}) to report with value: {text}")

        if data_format == "numbered_list":
            logging.debug("Formatting a numbered list")
            if text:
                for list_idx, list_entry in enumerate(text):
                    self._write_text(para, f"{list_idx+1}. {list_entry}\n")
            else:
                logging.warning(
                    f"Empty list has been found: {text}.  Resetting value to null string"
                )
                self._write_text(para, "")
        elif data_format == "test_result":
            logging.debug("Formatting a Test Result")
            if text:
                self._write_text(para, "PASS", bold=True, color=RGBColor(0, 255, 0))
            else:
                self._write_text(para, "FAIL", bold=True, color=RGBColor(255, 0, 0))
        else:
            self._write_text(
                para, text, font=font, font_size=font_size, bold=bold, color=text_color
            )

        if color:
            cell = table.cell(row, column)

            color = parse_xml(
                # pylint: disable-next=consider-using-f-string
                r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color)
            )
            # pylint: disable-next=protected-access
            cell._tc.get_or_add_tcPr().append(color)

    # pylint: disable-next=too-many-arguments
    def _write_text(
        self,
        para,
        text,
        font=None,
        font_size=None,
        bold=False,
        align=None,
        color=None,
        left_indent=None,
    ):
        """Writes a cell within Word doc table

        Args:
            para (obj): Word doc obj representing a paragraph
            text (str): Text to output in table cell
            font (str, optional): Font to use in table cell. Defaults to None.
            font_size (int, optional): Font size to use in table cell. Defaults to None.
            bold (bool, optional): Bold text in table cell. Defaults to False.
            align (obj, optional): Set paragraph alignment.  Defaults to None
            color (obj, optional): Text output color. Defaults to None
            left_indent (obj, optional): Sets left in indent in Doc. Defaults to None
        """
        run = para.add_run(text)
        logging.debug(
            f"Adding text '{text}' with font {font}, font_size {font_size}, bold {bold}, "
            f"alignment {align}"
        )

        if font:
            run.font.name = font

        if font_size:
            run.font.size = Pt(font_size)

        if bold:
            run.font.bold = bold

        if align == "right":
            # pylint: disable-next=no-member
            para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "center":
            # pylint: disable-next=no-member
            para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        if color:
            run.font.color.rgb = color

        if left_indent:
            para.paragraph_format.left_indent = left_indent

    def _return_summary_headers(self, report_template):
        """Parses report_table for summary fields and returns them

        Args:
            report_template (dict): Data structure describing reports fields

        Returns:
            list: List of fields to include in test case report header row
        """
        logging.debug(f"Checking report style for summary headers {report_template}")
        summary_headers = {k: v for (k, v) in report_template.items() if "summary" in v}
        logging.debug(f"The following are summary fields: {summary_headers}")
        return summary_headers

    def _compile_custom_tc_results(self, summary_headers):
        """Find subset of test cases using summary headers

        Args:
            summary_headers (list): List of fields to include in test case report header row

        Returns:
            dict: Data structure with test case results
        """
        if not self._results_datamodel:
            logging.warning("Skipping test case results")
            return None

        tbl_headers = summary_headers.keys()
        test_suites = self._results_datamodel["test_suites"]
        testcase_results = []

        for test_suite in test_suites:
            test_cases = test_suite["test_cases"]

            for test_case in test_cases:
                tc_name = self._format_tc_name(test_case["name"])
                logging.info(f"Compiling results for test case {tc_name}")
                duts = test_case["duts"]

                for dut in duts:
                    testcase_result = {}

                    for tbl_header in tbl_headers:
                        tbl_value = self._return_tbl_value(dut, tbl_header)
                        testcase_result[tbl_header] = tbl_value

                    logging.debug(f"Compiled DUT results: {testcase_result}")
                    testcase_results.append(testcase_result)

        logging.info("Returning testcase result")
        logging.debug(f"Returning testcase result {testcase_results}")
        return testcase_results

    def _return_tbl_value(self, dut, tbl_header):
        """Return a test case value for a summary header

        Args:
            dut (dict): dictionary of test parameters
            tbl_header (str): Summary header field

        Returns:
            str: Test case value for summary header
        """
        logging.debug(f"dut data structure set to: {dut}")
        if tbl_header in dut:
            tbl_value = dut[tbl_header]
            logging.debug(f"{tbl_header} set to {tbl_value} in dut structure")
        else:
            logging.warning(f"{tbl_header} NOT in dut structure")
            tbl_value = "Value NOT set in test case"

        return tbl_value

    def _default_tc_report(self):
        """Write default summary test case report"""

        testcase_results = self._compile_testcase_results()
        if not testcase_results:
            logging.warning("Skipping the summary testcase report")
            return

        table = self._document.add_table(rows=1, cols=7)
        table.style = TABLE_GRID
        test_num = 1

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Serial No"
        hdr_cells[1].text = "Test Id"
        hdr_cells[2].text = "Test Suite"
        hdr_cells[3].text = "Test Case"
        hdr_cells[4].text = "DUT/s"
        hdr_cells[5].text = "Result"
        hdr_cells[6].text = "Failure or Skip Reason"
        if not testcase_results:
            logging.warning("Skipping the summary testcase report")
            return
        for testcase_result in testcase_results:
            row_cells = table.add_row().cells
            row_cells[0].text = str(test_num)
            row_cells[1].text = str(testcase_result["test_id"])
            row_cells[2].text = str(testcase_result["test_suite"])
            row_cells[3].text = str(testcase_result["test_case"])
            row_cells[4].text = str(testcase_result["dut"])
            row_cells[5].text = str(testcase_result["results"])
            row_cells[6].text = str(testcase_result["fail_or_skip_reason"])

            test_num += 1

        self._document.add_page_break()

    def _write_detail_report(self):
        """Write detailed test case report"""

        logging.info("Creating report detailed test case report section")
        if not self._results_datamodel:
            logging.warning("Skipping the detailed testcase report")
            return

        test_suites = self._results_datamodel["test_suites"]

        for test_suite in test_suites:
            self._write_detail_major_section(test_suite)
            minor_section = 1

            for test_case in test_suite["test_cases"]:
                self._write_detail_minor_section(test_case, minor_section)
                dut_section = 1

                for dut in test_case["duts"]:
                    self._write_detail_dut_section(dut, minor_section, dut_section)
                    dut_section += 1

                minor_section += 1

            self._major_section += 1

    def _write_detail_major_section(self, test_suite):
        """Write detailed major report section

        Args:
            test_case (dict): Test case result specific parameters
        """

        ts_name = self._format_ts_name(test_suite["name"])
        self._document.add_heading(
            f"{self._major_section}. Detailed Test Suite Results: {ts_name}",
            1,
        )

    def _write_detail_minor_section(self, test_case, minor_section):
        """Write detailed minor report section

        Args:
            test_case (dict): Test case result specific parameters
            minor_section (int): minor section number
        """

        tc_name = self._format_tc_name(test_case["name"])

        self._document.add_heading(f"{self._major_section}.{minor_section} Test Case: {tc_name}", 2)

    def _write_detail_dut_section(self, dut, minor_section, dut_section):
        """Write detailed DUT test case report section

        Args:
            dut (dict): DUT specific parameters
            minor_section (int): Minor section number
            dut_section (int): DUT section number
        """

        logging.debug(f"Raw DUT data is {dut}")
        dut_name = dut["dut"]
        dut_name = dut_name.upper()
        tc_name = dut["name"]
        logging.info(f"Writing detailed test case results: {tc_name} for dut: {dut_name}")
        logging.debug(f"DUT name is {dut_name}")
        self._document.add_heading(
            f"{self._major_section}.{minor_section}. {dut_section} DUT: {dut_name}",
            3,
        )

        report_styles = REPORT_TEMPLATES.keys()
        logging.debug(f"Inputted the following report styles {report_styles}")

        if "report_style" in dut:
            if dut["report_style"] in report_styles:
                logging.debug(f"Report_style is set: {dut['report_style']}, custom")
                self._write_custom_detail_dut_section(dut)
            else:
                logging.debug(f"Report_style is incorrect: {dut['report_style']}. Set to default.")
                dut["report_style"] = "default"
                self._write_custom_detail_dut_section(dut)
        else:
            logging.warning("Report_style is NOT set.  Setting to default report_style.")
            dut["report_style"] = "default"
            self._write_custom_detail_dut_section(dut)

    def _write_custom_detail_dut_section(self, dut):
        """Method for writing a custom detailed DUT section within report
            Uses report_template data structure to describe reporting

        Args:
            dut (dict): Data structure with DUT specific data
        """

        report_style = dut["report_style"]
        tc_name = dut["name"]
        report_template = REPORT_TEMPLATES[report_style]

        missing_fields = self._required_template_fields(dut, report_template)

        if missing_fields:
            logging.warning(
                "Required report fields are NOT in test definitions for test case: {tc_name}"
            )
            para = self._document.add_paragraph()
            out_msg = (
                f"The following required fields are missing: {missing_fields}, ",
                "Please update test definition file with this information ",
                "so test case report can be generated.",
            )
            self._write_text(para, out_msg, bold=True, color=RGBColor(255, 0, 0))
        else:
            logging.info(f"Required report fields are in test definitions for test case: {tc_name}")
            self._write_custom_paragraph(dut, report_template)

    def _required_template_fields(self, dut, report_template):
        """Uses report template to verify dut data structure has required
            fields

        Args:
            dut (dict): Data structure with DUT specific data
            report_template (dict): Data structure describing reports fields

        Returns:
            set: Required fields missing from dut data struct
        """
        required_fields = {k for (k, v) in report_template.items() if v["required"]}
        logging.warning(f"The following fields are required: {required_fields}")
        logging.debug(f"DUT keys are {dut.keys()}")
        missing_fields = required_fields.difference(dut.keys())
        logging.error(f"The following required fields are missing: {missing_fields}")

        return missing_fields

    # pylint: disable-next=too-many-branches
    def _write_custom_paragraph(self, dut, report_template):
        """Writes section of detailed dut report based on field's format

        Args:
            dut (dict): Data structure with DUT specific data
            report_template (dict): Data structure describing reports fields
        """

        for report_field in report_template:
            para = self._write_output_name(report_template, report_field)

            if report_field not in dut:
                self._set_default_value(report_field, report_template, dut)

            if "format" in report_template[report_field]:
                report_format = report_template[report_field]["format"]
                if report_field in dut:
                    report_text = dut[report_field]
                else:
                    report_text = "Field NOT set in test case"
            else:
                report_format = "missing"

            logging.debug(f"Format for {report_field} is set to {report_format}")

            if report_format == "string":
                self._write_text(para, report_text)
            elif report_format == "bulleted_list":
                self._write_bulleted_list(dut, report_field)
            elif report_format == "numbered_list":
                self._write_numbered_list(dut, report_field)
            elif report_format == "dict_string":
                self._write_dict_string(dut, report_field)
            elif report_format == "config_string":
                self._write_config_string(dut, report_field)
            elif report_format == "config_term":
                self._write_config_term(dut, report_field)
            elif report_format == "test_result":
                self._write_test_result(dut, para, report_field)
            else:
                run = para.add_run(
                    "Please correctly set format in report template for "
                    f"{report_field} field. Output cannot be displayed "
                    "without this."
                )
                run.font.color.rgb = RGBColor(255, 0, 0)

    def _set_default_value(self, report_field, report_template, dut):
        """Uses default value in report template and sets it in dut

        Args:
            report_field (dict): Data Struct describing a reports field
            dut (dict): Data structure with DUT specific data
            report_template (dict): Data structure describing reports fields
        """
        default_value = report_template[report_field]["default"]
        dut[report_field] = default_value

    def _write_output_name(self, report_template, report_field):
        """Writes an output section in detailed dut section

        Args:
            report_template (dict): Data structure describing reports fields
            report_field (string): Name of report field in dut to write

        Returns:
            obj: Current word paragraph object
        """
        para = self._document.add_paragraph()
        if "output_name" in report_template[report_field]:
            output_name = report_template[report_field]["output_name"]
        else:
            output_name = report_field

        self._write_text(para, f"{output_name.upper()}: ", bold=True)

        return para

    def _write_bulleted_list(self, dut, report_field):
        """Write a generic bulleted list to Word doc

        Args:
            dut (dict): Data structure with DUT specific data
            report_field (string): Name of report field in dut to write
        """

        if report_field in dut:
            report_value = dut[report_field]
            para = self._document.add_paragraph()
            para.style = "List Bullet 2"
            self._write_text(para, report_value)

    def _write_numbered_list(self, dut, report_field):
        """Write a generic numbered list to Word doc

        Args:
            dut (dict): Data structure with DUT specific data
            report_field (string): Name of report field in dut to write
        """

        if report_field in dut:
            report_values = dut[report_field]
            list_counter = 1

            for report_value in report_values:
                report_value = f"{list_counter}. {report_value}"
                para = self._document.add_paragraph()
                self._write_text(para, report_value, left_indent=Inches(0.25))
                list_counter += 1

    def _write_dict_string(self, dut, report_field):
        """Write a dictionary in YAML format to Word doc

        Args:
            dut (dict): Data structure with DUT specific data
            report_field (string): Name of report field in dut to write
        """

        if report_field in dut:
            report_value = dut[report_field]
            formatted_data = yaml.dump(report_value)
            logging.debug(f"Data formatted to YAML: {formatted_data}")
            para = self._document.add_paragraph()
            self._write_text(para, formatted_data.strip(), left_indent=Inches(0.25))

    def _write_test_result(self, dut, para, report_field):
        """Write test result string to Word doc with formatting

        Args:
            dut (dict): Data structure with DUT specific data
            para (obj): Current Word paragraph object
            report_field (string): Name of report field in dut to write
        """

        if report_field in dut:
            report_value = dut[report_field]
            if report_value:
                self._write_text(para, "PASS", bold=True, color=RGBColor(0, 255, 0))
            else:
                self._write_text(para, "FAIL", bold=True, color=RGBColor(255, 0, 0))

    def _write_config_string(self, dut, report_field):
        """Write list of EOS configurations to Word doc with formatting

        Args:
            dut (dict): Data structure with DUT specific data
            report_field (string): Name of report field in dut to write
        """

        if report_field in dut:
            report_value = dut[report_field]
            para = self._document.add_paragraph()
            self._write_text(
                para,
                report_value,
                font="Courier New",
                font_size=10,
                left_indent=Inches(0.25),
            )

    def _write_config_term(self, dut, report_field):
        """Write EOS configurations to Word doc with formatting to appear
           from a xterm

        Args:
            dut (dict): Data structure with DUT specific data
            report_field (string): Name of report field in dut to write
        """

        if report_field in dut and "show_cmds" in dut:
            table = self._document.add_table(rows=1, cols=1, style="Table Grid")

            show_cmd_txts = dut["show_cmd_txts"]
            show_cmds = dut["show_cmds"]
            index = 0
            for dut_name in show_cmds.keys():
                for command, text in zip(show_cmds.get(dut_name), show_cmd_txts.get(dut_name)):
                    if index != 0:
                        _ = table.add_row().cells
                    config_output = f"\n{dut_name}# {command}\n\n{text}\n"
                    self._write_cell(
                        table,
                        config_output,
                        0,
                        index,
                        font="Courier New",
                        font_size=10,
                        color="0A0A0A",
                        text_color=RGBColor(0, 255, 0),
                    )
                    index = index + 1

            para = self._document.add_paragraph()
            _ = para.add_run()

    def _compile_suite_results(self):
        """Compile test suite results and return them

        Return:
            suite_results (list): List of compiled test suite data
        """

        logging.debug(f"The following test suites have been collected {self._results_datamodel}")

        if not self._results_datamodel:
            logging.warning("Skipping the compiled test suite result")
            return ""

        test_suites = self._results_datamodel["test_suites"]
        suite_results = []

        for test_suite in test_suites:
            suite_result = {}
            suite_result["total_tests"] = 0
            suite_result["total_pass"] = 0
            suite_result["total_fail"] = 0
            suite_result["total_skip"] = 0
            suite_result["name"] = test_suite["name"]

            suite_name = suite_result["name"]
            logging.warning(
                "Zeroing test_suite results for test suite: "
                f"{suite_name} and data: {suite_result}"
            )

            for test_case in test_suite["test_cases"]:
                for dut in test_case["duts"]:
                    suite_result["total_tests"] += 1

                    if dut["test_result"] and dut["test_result"] == "Skipped":
                        suite_result["total_skip"] += 1
                        dut["fail_or_skip_reason"] = dut.get("actual_output", "")
                    elif dut["test_result"] and dut["test_result"] != "Skipped":
                        suite_result["total_pass"] += 1
                    else:
                        suite_result["total_fail"] += 1

            logging.debug(f"Compiled test suite data: {suite_result}")
            suite_results.append(suite_result)

        logging.debug(f"Compiled suite results: {suite_results}")
        return suite_results

    def _compile_testcase_results(self):
        """Compile test case results and return them"""

        if not self._results_datamodel:
            logging.warning("Skipping test case results")
            return ""

        test_suites = self._results_datamodel["test_suites"]
        testcase_results = []

        for test_suite in test_suites:
            test_cases = test_suite["test_cases"]
            ts_name = self._format_ts_name(test_suite["name"])
            logging.info(f"Compiling results for test suite {ts_name}")

            for test_case in test_cases:
                tc_name = self._format_tc_name(test_case["name"])
                logging.info(f"Compiling results for test case {tc_name}")
                duts = test_case["duts"]

                for dut in duts:
                    testcase_result = {}

                    dut_name = dut["dut"]
                    fail_reason = dut["fail_or_skip_reason"]
                    logging.debug(f"Compiling results for DUT/s {dut_name}")
                    testcase_id = dut["test_id"]

                    if dut["test_result"] and dut["test_result"] == "Skipped":
                        test_result = "SKIP"
                        fail_reason = dut.get("actual_output", "")
                    elif dut["test_result"] and dut["test_result"] != "Skipped":
                        test_result = "PASS"
                    else:
                        test_result = "FAIL"

                    testcase_result["test_suite"] = ts_name
                    testcase_result["test_case"] = tc_name
                    testcase_result["test_id"] = testcase_id
                    testcase_result["dut"] = dut_name
                    testcase_result["results"] = test_result
                    testcase_result["fail_or_skip_reason"] = fail_reason
                    logging.debug(f"Compiled results: {testcase_result}")

                    testcase_results.append(testcase_result)
                    logging.debug(f"After testcase results struct appended: {testcase_results}")

                logging.debug(f"Interim dut -- testcase results struct {testcase_results}")

        logging.debug(f"Returning testcase result {testcase_results}")
        return testcase_results

    def _format_ts_name(self, ts_name):
        """Input a test suite program name and return a formatted name for
            test suite

        Args:
            ts_name (str): Name of test suite program

        Return:
            ts_name (str): Formatted test suite name
        """

        logging.debug(f"Test suite name is {ts_name}")
        ts_name = ts_name.split(".")[0]
        ts_name = ts_name.split("_")
        if len(ts_name) > 1:
            ts_name = ts_name[1].capitalize()
        logging.debug(f"Formatted test suite name is {ts_name}")

        return ts_name

    def _format_tc_name(self, tc_name):
        """Input a PyTest test case name and return a formatted name for
            test case

        Args:
            tc_name (str): Name of PyTest test case

        Return:
            tc_name (str): Formatted test case name
        """

        logging.debug(f"Test case name is {tc_name}")
        tc_name = tc_name.replace("_", " ")
        tc_name = tc_name.replace("intf", "interface")
        tc_name = tc_name.capitalize()

        logging.debug(f"Formattted test case name is {tc_name}")

        return tc_name
