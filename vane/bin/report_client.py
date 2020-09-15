#!/usr/bin/python3
#
# Copyright (c) 2019, Arista Networks EOS+
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# * Redistributions of source code must retain the above copyright notice, this
#   list of conditions and the following disclaimer.
#
# * Redistributions in binary form must reproduce the above copyright notice,
#   this list of conditions and the following disclaimer in the documentation
#   and/or other materials provided with the distribution.
#
# * Neither the name of the Arista nor the names of its
#   contributors may be used to endorse or promote products derived from
#   this software without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE
# FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL
# DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
# SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER
# CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
# OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
# OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#

"""Utilities for using PyTest in network testing"""

import concurrent.futures
import time
import fcntl
import sys
import logging
import os
import yaml
import docx
import datetime
import json
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH


logging.basicConfig(level=logging.INFO, filename='vane.log', filemode='w',
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')



class ReportClient:
    """ Creates an instance of the Report Client.
    """

    def __init__(self, test_definition):
        """ Initializes the Test Client

            Args:
                test_definition (str): YAML representation of NRFU tests
        """

        logging.info('Convert yaml data-model to a python data structure')
        self.data_model = self._import_yaml(test_definition)
        logging.info('Internal test data-model initialized with value: '
                     f'{self.data_model}')
        
        _results_file = self.data_model['parameters']['results_file']
        self._results_data = self._import_yaml(_results_file)
        logging.info(f'Results file data is {self._results_data}')

        self._document = docx.Document()
        self._summary_results = self._compile_test_results()
        logging.info(f"Test Results: {self._summary_results}")
        self._major_section = 1

    def _import_yaml(self, yaml_file):
        """ Import YAML file as python data structure

            Args:
                yaml_file (str): Name of YAML file
        """

        logging.info(f'Opening {yaml_file} for read')
        try:
            with open(yaml_file, 'r') as input_yaml:
                try:
                    yaml_data = yaml.safe_load(input_yaml)
                    logging.info(f'Inputed the following yaml: '
                                 f'{yaml_data}')
                    return yaml_data
                except yaml.YAMLError as err_data:
                    logging.error(f'Error in YAML file. {err_data}')
                    sys.exit(1)
        except OSError as err_data:
            logging.error(f'Defintions file: {yaml_file} not '
                          f'found. {err_data}')
            sys.exit(1)

    def write_result_doc(self):
        """ Create MSFT docx with results
        """

        logging.info('Create MSFT docx with results')
        self._write_title_page()
        self._write_summary_report()
        self._write_detail_report()
        self._document.save('../reports/report-class.docx')

    def _return_date(self):
        """ Genreate a formatted date and return to calling
            function.
        """

        date_obj = datetime.datetime.now()
        format_date = date_obj.strftime("%B %d, %Y %I:%M:%S%p")

        logging.info(f'Returning formatted date: {format_date}')
        return format_date

    def _write_title_page(self):
        """ Write report title page
        """

        logging.info('Create report title page')
        format_date = self._return_date()
        self._document.add_heading('Test Report', 0)
        p = self._document.add_paragraph(f'{format_date}')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._document.add_page_break()

    def _write_summary_report(self):
        """ Write summary reports
        """

        self._document.add_heading(f'{self._major_section}. Test Results Summary', 1)
        self._write_summary_results()
        self._write_dut_summary_results()
        self._write_suite_summary_results()

        self._major_section += 1
        self._document.add_page_break()

    def _write_summary_results(self):
        """ Write summary results section
        """

        logging.info("Create summary results table")
        self._document.add_heading(f'{self._major_section}.1 Summary Results', 2)
        table = self._document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Total Tests'
        hdr_cells[1].text = 'Total Passed'
        hdr_cells[2].text = 'Total Failed'
        hdr_cells[3].text = 'Total Skipped'
        hdr_cells[4].text = 'Total Errored'
        hdr_cells[5].text = 'Total Duration'

        ptr = self._summary_results['summaryResults']
        total_tests = self._totals(ptr, 'num_tests')
        total_pass = self._totals(ptr, 'passed')
        total_fail = self._totals(ptr, 'failed')
        total_skip = self._totals(ptr, 'skipped')
        total_err = self._totals(ptr, 'error')
        total_time = self._totals(ptr, 'duration')

        row_cells = table.add_row().cells
        row_cells[0].text = total_tests
        row_cells[1].text = total_pass
        row_cells[2].text = total_fail
        row_cells[3].text = total_skip
        row_cells[4].text = total_err
        row_cells[5].text = total_time

    def _write_dut_summary_results(self):
        """ Write summary DUT result section
        """

        logging.info("Create DUT summary results table")
        self._document.add_heading(f'{self._major_section }.2 Summary Totals for Devices Under Tests', 2)

        table = self._document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'DUT'
        hdr_cells[1].text = 'Total Tests'
        hdr_cells[2].text = 'Total Passed'
        hdr_cells[3].text = 'Total Failed'
        hdr_cells[4].text = 'Total Skipped'
        hdr_cells[5].text = 'Total Errored'

        duts = self._summary_results['duts']

        for dut in duts:
            total_tests = self._totals(dut, 'TOTAL')
            total_pass = self._totals(dut, 'PASS')
            total_fail = self._totals(dut, 'FAIL')
            total_skip = self._totals(dut, 'SKIP')
            total_err = self._totals(dut, 'ERROR')
            dut_name = self._totals(dut, 'name')

            row_cells = table.add_row().cells
            row_cells[0].text = dut_name
            row_cells[1].text = total_tests
            row_cells[2].text = total_pass
            row_cells[3].text = total_fail
            row_cells[4].text = total_skip
            row_cells[5].text = total_err

    def _write_suite_summary_results(self):
        """ Write summary test suite result section
        """

        # TODO: Use conftest get_closest_marker() to find markers
        logging.info("Create Mark summary results table")
        self._document.add_heading(f'{self._major_section }.3 Summary Totals for Test Suites', 2)

    def _compile_test_results(self):
        """ Parse PyTest JSON results and compile:
        """

        json_report = self.data_model['parameters']['json_report']
        json_report = f"{json_report}.json"
        test_results = {}
        logging.info(f'Opening JSON file {json_report} to parse for summary results')

        with open(json_report, 'r') as json_file:
            logging.info(f'Raw json report is {json_file}')
            test_data = json.load(json_file)
            logging.info(f'Structured json report is {test_data}')
            test_results["summaryResults"] = test_data["report"]["summary"]
            logging.info(f'Summary for test cases are {test_results["summaryResults"]}')
            test_results["duts"] = self._parse_testcases(test_data["report"]["tests"])
        
        return test_results

    def _parse_testcases(self, testcases):
        """ Parse Test cases and return compilation per DUT
        """

        testcases_results = []
        dut_list = []

        for testcase in testcases:
            if re.search('\[.*\]', testcase["name"]):
                dut_name = re.findall('\[.*\]', testcase["name"])[0][1:-1]
                test_result = testcase["outcome"]

                if dut_name not in dut_list:
                    dut_list.append(dut_name)
                    testcases_results.append({})
                    testcases_results[-1]["PASS"] = 0
                    testcases_results[-1]["FAIL"] = 0
                    testcases_results[-1]["SKIP"] = 0
                    testcases_results[-1]["ERROR"] = 0
                    testcases_results[-1]["TOTAL"] = 0

                dut_index = dut_list.index(dut_name)
                testcases_results[dut_index]["name"] = dut_name

                if test_result == "passed":
                    testcases_results[dut_index]["PASS"] += 1
                elif test_result == "failed":
                    testcases_results[dut_index]["FAIL"] += 1
                elif test_result == "skipped":
                    testcases_results[dut_index]["SKIP"] += 1
                elif test_result == "error":
                    testcases_results[dut_index]["ERROR"] += 1

                testcases_results[dut_index]["TOTAL"] += 1

        return testcases_results

    def _totals(self, ptr, ptr_key):
        """ Test for a key in dictionary.  If key exists return key and if key is
            missing return 0

        Args:
            ptr (dict): dictionary to check
            ptr_key (str): key to test if in dict
        """

        if ptr_key in ptr:
            return str(ptr[ptr_key])
        else:
            return "0"
    
    def _write_detail_report(self):
        """ Write summary reports
        """

        test_suites = self._results_data['test_suites']

        for test_suite in test_suites:
            self._write_detail_major_section(test_suite)
            minor_section = 1

            for test_case in test_suite['test_cases']:
                self._write_detail_minor_section(test_case, minor_section)
                dut_section = 1

                for dut in test_case['duts']:
                    self._write_dut_minor_section(dut, minor_section, dut_section)
                    dut_section += 1

                minor_section += 1

            self._major_section += 1

    def _write_detail_major_section(self, test_suite):
        """ Write Detailed majore report section

            Args:
                test_suite (dict): test_suite result data
        """

        logging.info(f'Raw test suite data is {test_suite}')
        ts_name = test_suite['name']
        logging.info(f'Test suite name is {ts_name}')
        ts_name = ts_name.split('.')[0]
        ts_name = ts_name.split('_')[1].upper()
        logging.info(f'Formatted test suite name is {ts_name}')
        self._document.add_heading(f'{self._major_section}. Detailed Test '
                                   f'Suite Results: {ts_name}', 1)
    
    def _write_detail_minor_section(self, test_case, minor_section):
        """[summary]

        Args:
            test_case (dict): test_case result data
            minor_section (int): minor section number
        """

        logging.info(f'Raw test case data is {test_case}')
        tc_name = test_case['name']
        logging.info(f'Test case name is {tc_name}')
        tc_name = ' '.join(tc_name.split('_'))[:-3].upper()
        logging.info(f'Formattted test case name is {tc_name}')
        self._document.add_heading(f'{self._major_section}.{minor_section} Test Case: {tc_name}', 2)
        description = test_case['duts'][0]['description']
        p = self._document.add_paragraph(f'Description: {description}')

    def _write_dut_minor_section(self, dut, minor_section, dut_section):
        """[summary]
        Args:
            dut ([type]): [description]
            minor_section ([type]): [description]
            dut_section ([type]): [description]
        """

        logging.info(f'Raw DUT data is {dut}')
        dut_name = dut['dut']
        dut_name = dut_name.upper()
        logging.info(f'DUT name is {dut_name}')
        self._document.add_heading(f'{self._major_section}.{minor_section}.{dut_section} DUT: {dut_name}', 3)


def write_result_doc():
    """ Create MSFT docx with results
    """

    # TODO: remove hard code
    yaml_file = 'result.yml'
    yaml_data = yaml_io(yaml_file, 'read')

    date_obj = datetime.datetime.now()
    format_date = date_obj.strftime("%B %d, %Y %I:%M:%S%p")

    logging.info('Create MSFT docx with results')
    document = docx.Document()

    logging.info('Create report title page')
    document.add_heading('Test Report', 0)
    p = document.add_paragraph(f'{format_date}')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_page_break()

    document.add_heading(f'1. Test Results Summary', 1)
    summary_results = compile_test_results()
    logging.info(f"Test Results: {summary_results}")

    logging.info("Create summary results table")
    document.add_heading(f'1.1 Summary Results', 2)
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Total Tests'
    hdr_cells[1].text = 'Total Passed'
    hdr_cells[2].text = 'Total Failed'
    hdr_cells[3].text = 'Total Skipped'
    hdr_cells[4].text = 'Total Errored'
    hdr_cells[5].text = 'Total Duration'

    ptr = summary_results['summaryResults']
    total_tests = totals(ptr, 'num_tests')
    total_pass = totals(ptr, 'passed')
    total_fail = totals(ptr, 'failed')
    total_skip = totals(ptr, 'skipped')
    total_err = totals(ptr, 'error')
    total_time = totals(ptr, 'duration')

    row_cells = table.add_row().cells
    row_cells[0].text = total_tests
    row_cells[1].text = total_pass
    row_cells[2].text = total_fail
    row_cells[3].text = total_skip
    row_cells[4].text = total_err
    row_cells[5].text = total_time

    document.add_heading(f'1.2 Summary Totals Devices Under Tests', 2)

    document.add_page_break()

    major_section = 2
    for test_suite in yaml_data['test_suites']:
        logging.info(f'Raw test suite data is {test_suite}')

        ts_name = test_suite['name']
        logging.info(f'Test suite name is {ts_name}')

        ts_name = ts_name.split('.')[0]
        ts_name = ts_name.split('_')[1].upper()
        logging.info(f'Formatted test suite name is {ts_name}')
        document.add_heading(f'{major_section}. Test Suite: {ts_name}', 1)

        minor_section = 1

        for test_case in test_suite['test_cases']:
            logging.info(f'Raw test case data is {test_case}')
            tc_name = test_case['name']
            logging.info(f'Test case name is {tc_name}')
            tc_name = ' '.join(tc_name.split('_'))[:-3].upper()
            logging.info(f'Formattted test case name is {tc_name}')
            document.add_heading(f'{major_section}.{minor_section} Test Case: {tc_name}', 2)

            description = test_case['duts'][0]['description']
            p = document.add_paragraph(f'DESCRIPTION: {description}')

            maintenance_section = 1

            for dut in test_case['duts']:
                logging.info(f'Raw DUT data is {dut}')
                dut_name = dut['dut']
                dut_name = dut_name.upper()
                logging.info(f'DUT name is {dut_name}')
                document.add_heading(f'{major_section}.{minor_section}.{maintenance_section} DUT: {dut_name}', 3)

                maintenance_section += 1

            minor_section += 1

        major_section += 1


    """
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()
    """

    document.save('../reports/report.docx')


def yaml_io(yaml_file, io_type, yaml_data=None):
    """ Write test results to YAML file for post-processing

        Args:
            yaml_file (str): Name of YAML file
            io (str): Read or write to YAML file
    """

    while True:
        try:
            if io_type == 'read':
                with open(yaml_file, 'r') as yaml_in:
                    fcntl.flock(yaml_in, fcntl.LOCK_EX | fcntl.LOCK_NB)
                    yaml_data = yaml.safe_load(yaml_in)
                    fcntl.flock(yaml_in, fcntl.LOCK_UN)
                    break
            else:
                with open(yaml_file, 'w') as yaml_out:
                    fcntl.flock(yaml_out, fcntl.LOCK_EX | fcntl.LOCK_NB)
                    yaml.dump(yaml_data, yaml_out, default_flow_style=False)
                    fcntl.flock(yaml_out, fcntl.LOCK_UN)
                    break
        except:
            time.sleep(0.05)

    return yaml_data


def totals(ptr, ptr_key):
    """ Test for a key in dictionary.  If key exists return key and if key is
        missing return 0

    Args:
        ptr (dict): dictionary to check
        ptr_key (str): key to test if in dict
    """

    if ptr_key in ptr:
        return str(ptr[ptr_key])
    else:
        return "0"


def compile_test_results():
    """ Parse NRFU results and compile:
        test_results: <Pass/Fail>
        Pass: <number passed>
        Fail: <number failed>
        Duts:
          - test_results: <Pass/Fail>
            Pass: <number passed>
            FAil: <number failed>
    """

    # TODO: Remove hard code
    json_report = "../reports/report.json"
    test_results = {}

    while True:
        try:
            with open(json_report, 'r') as json_file:
                fcntl.flock(json_file, fcntl.LOCK_EX | fcntl.LOCK_NB)
                test_data = json.load(json_file)
                fcntl.flock(json_file, fcntl.LOCK_UN)

                test_results["summaryResults"] = test_data["report"]["summary"]
                logging.info(f'Summary for test cases are {test_results["summaryResults"]}')
                test_results["duts"] = parse_testcases(test_data["report"]["tests"])
                break

        except:
            time.sleep(0.05)

    return test_results


def parse_testcases(testcases):
    """ Parse Test cases and return compilation per DUT
    """

    testcases_results = []
    dut_list = []

    for testcase in testcases:
        if re.search('\[.*\]', testcase["name"]):
            dut_name = re.findall('\[.*\]', testcase["name"])[0][1:-1]
            test_result = testcase["outcome"]

            if dut_name not in dut_list:
                dut_list.append(dut_name)
                testcases_results.append({})
                testcases_results[-1]["PASS"] = 0
                testcases_results[-1]["FAIL"] = 0
                testcases_results[-1]["SKIP"] = 0
                testcases_results[-1]["ERROR"] = 0
                testcases_results[-1]["TOTAL"] = 0

            dut_index = dut_list.index(dut_name)
            testcases_results[dut_index]["name"] = dut_name

            if test_result == "passed":
                testcases_results[dut_index]["PASS"] += 1
            elif test_result == "failed":
                testcases_results[dut_index]["FAIL"] += 1
            elif test_result == "skipped":
                testcases_results[dut_index]["SKIP"] += 1
            elif test_result == "error":
                testcases_results[dut_index]["ERROR"] += 1
            
            testcases_results[dut_index]["TOTAL"] += 1

    return testcases_results